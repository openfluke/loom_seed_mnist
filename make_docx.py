#!/usr/bin/env python3
"""Pack loom_seed_mnist + loom/poly seed stack into DOCX files.

Writes two documents by default (seed stack first):
  1) docs/poly_seed_*.docx     — loom/poly seed_* (+ dense helpers the PoC uses)
  2) docs/seed_mnist_*.docx    — this repo README / Go / run.sh
"""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
LOOM = ROOT.parent / "loom"
POLY = LOOM / "poly"
DOCS = ROOT / "docs"

# This PoC's own sources
PROJECT_GLOBS = [
    "README.md",
    "main.go",
    "go.mod",
    "run.sh",
    "seedmnist/*.go",
]

# Seed system loom_seed_mnist sits on (+ dense builder it needs)
POLY_SEED_NAMES = [
    "seed_core.go",
    "seed_init.go",
    "seed_dtypes.go",
    "seed_dtypes_layers.go",
    "seed_infinite.go",
    "seed_infinite_dtypes.go",
    "seed_manifest.go",
    "seed_entity.go",
    "seed_invert.go",
    "seed_dense.go",
    "seed_swiglu.go",
    "seed_mha.go",
    "seed_cnn.go",
    "seed_rnn.go",
    "seed_lstm.go",
    "seed_embedding.go",
    "seed_residual.go",
    # Dense build path used when materializing nets from seeds
    "dense.go",
    "architecture.go",
]


def collect_files(globs: list[str], base: Path = ROOT) -> list[Path]:
    found: list[Path] = []
    seen: set[Path] = set()
    for pattern in globs:
        for p in sorted(base.glob(pattern)):
            if p.is_file() and p.resolve() not in seen:
                seen.add(p.resolve())
                found.append(p)
    return found


def collect_poly_seed_files() -> list[Path]:
    out: list[Path] = []
    for name in POLY_SEED_NAMES:
        p = POLY / name
        if p.is_file():
            out.append(p)
    return out


def set_run_font(run, *, mono: bool = False, size: int = 10) -> None:
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    if mono:
        run.font.name = "Courier New"
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), "Courier New")
        rFonts.set(qn("w:hAnsi"), "Courier New")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_code_block(doc: Document, text: str, *, max_chars: int | None = None) -> None:
    body = text
    truncated = False
    if max_chars is not None and len(body) > max_chars:
        body = body[:max_chars]
        truncated = True
    for line in body.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line if line else " ")
        set_run_font(run, mono=True, size=8)
    if truncated:
        p = doc.add_paragraph()
        run = p.add_run(f"\n… truncated at {max_chars} characters …")
        set_run_font(run, mono=True, size=8)


def add_file_section(
    doc: Document,
    path: Path,
    *,
    rel: Path | None = None,
    max_chars: int | None = None,
) -> None:
    label = str(rel) if rel is not None else path.name
    add_heading(doc, label, level=2)
    meta = doc.add_paragraph()
    run = meta.add_run(f"{path}  ·  {path.stat().st_size} bytes")
    set_run_font(run, size=9)
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except OSError as e:
        p = doc.add_paragraph()
        run = p.add_run(f"(could not read: {e})")
        set_run_font(run, size=9)
        return
    add_code_block(doc, text, max_chars=max_chars)


def rel_label(path: Path) -> Path:
    try:
        return path.relative_to(ROOT)
    except ValueError:
        return Path("..") / path.relative_to(ROOT.parent)


def build_poly_seed_docx(out: Path, *, max_chars: int | None = None) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    files = collect_poly_seed_files()
    missing = [n for n in POLY_SEED_NAMES if not (POLY / n).is_file()]

    add_heading(doc, "loom/poly — seed stack (for loom_seed_mnist)", level=0)
    intro = doc.add_paragraph()
    run = intro.add_run(
        f"Generated {datetime.now().isoformat(timespec='seconds')}\n"
        f"Poly root: {POLY}\n"
        "layer_seed → He-init PRNG → weight matrices. This is what loom_seed_mnist searches.\n"
        f"Files included: {len(files)}"
        + (f"  · missing: {', '.join(missing)}" if missing else "")
    )
    set_run_font(run, size=10)

    add_heading(doc, "File index", level=1)
    for p in files:
        try:
            rel = p.relative_to(LOOM)
        except ValueError:
            rel = Path(p.name)
        doc.add_paragraph(f"• {rel} ({p.stat().st_size} B)", style="List Bullet")

    add_heading(doc, "seed_* sources", level=1)
    for p in files:
        if not p.name.startswith("seed_"):
            continue
        try:
            rel = p.relative_to(LOOM)
        except ValueError:
            rel = Path(p.name)
        add_file_section(doc, p, rel=rel, max_chars=max_chars)

    helpers = [p for p in files if not p.name.startswith("seed_")]
    if helpers:
        add_heading(doc, "Dense / architecture helpers", level=1)
        for p in helpers:
            try:
                rel = p.relative_to(LOOM)
            except ValueError:
                rel = Path(p.name)
            add_file_section(doc, p, rel=rel, max_chars=max_chars)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def build_seed_mnist_docx(out: Path, *, max_chars: int | None = None) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    sources = collect_files(PROJECT_GLOBS)

    add_heading(doc, "loom_seed_mnist — README + code", level=0)
    intro = doc.add_paragraph()
    run = intro.add_run(
        f"Generated {datetime.now().isoformat(timespec='seconds')}\n"
        f"Root: {ROOT}\n"
        "PoC: search layer_seed values only (no free weight training).\n"
        "Companion poly seed stack is in the separate poly_seed_*.docx."
    )
    set_run_font(run, size=10)

    add_heading(doc, "File index", level=1)
    for p in sources:
        doc.add_paragraph(
            f"• {p.relative_to(ROOT)} ({p.stat().st_size} B)", style="List Bullet"
        )

    add_heading(doc, "README", level=1)
    readme = ROOT / "README.md"
    if readme.is_file():
        add_file_section(doc, readme, rel=Path("README.md"), max_chars=max_chars)
    else:
        doc.add_paragraph("(README.md missing)")

    add_heading(doc, "Source code", level=1)
    for p in sources:
        if p.name == "README.md":
            continue
        add_file_section(doc, p, rel=p.relative_to(ROOT), max_chars=max_chars)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def main() -> None:
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument(
        "--poly-seeds-output",
        type=Path,
        default=DOCS / f"poly_seed_{stamp}.docx",
        help="docx for loom/poly seed_* stack (written first)",
    )
    ap.add_argument(
        "-o",
        "--output",
        type=Path,
        default=DOCS / f"seed_mnist_{stamp}.docx",
        help="docx for this repo README + code",
    )
    ap.add_argument(
        "--poly-seeds-only",
        action="store_true",
        help="only write the poly_seed docx",
    )
    ap.add_argument(
        "--project-only",
        action="store_true",
        help="only write the seed_mnist project docx",
    )
    args = ap.parse_args()

    wrote: list[Path] = []
    # Seed stack first (as requested)
    if not args.project_only:
        wrote.append(build_poly_seed_docx(args.poly_seeds_output))
    if not args.poly_seeds_only:
        wrote.append(build_seed_mnist_docx(args.output))

    for path in wrote:
        print(f"wrote {path} ({path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
